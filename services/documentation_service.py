"""Generate engineering documentation packages from canonical tag data."""

from __future__ import annotations

import csv
import html
import re
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable, Iterable

from app_config import SYNC_STATUS_LABELS
from models.pending_export import PendingExportChange
from models.tag_record import (
    SYNC_NAME_MISMATCH,
    SYNC_NEEDS_ALIGN,
    SYNC_PROFICY_DRIFT,
    SYNC_PROFICY_ONLY,
    SYNC_SYNCED,
    TagRecord,
)
from services.address_normalizer import is_resolvable_address, normalize_address
from services.export_queue_service import changed_field_labels
from services.pandas_lazy import get_pandas

_ALARM_PATTERN = re.compile(
    r"\b(ALM|ALARM|ACKN|ACK|TRIP|FAULT|SHUTDOWN|INTERLOCK|HH|LL|HI|LO)\b",
    re.IGNORECASE,
)
_ADDRESS_FAMILY_PATTERN = re.compile(r"^%([A-Z]+)", re.IGNORECASE)

DOCUMENT_TYPES: dict[str, tuple[str, str]] = {
    "io_list": (
        "IO List",
        "PLC I/O points with addresses, descriptions, and program cross-references.",
    ),
    "alarm_list": (
        "Alarm List",
        "Alarm, trip, and interlock tags inferred from naming and descriptions.",
    ),
    "tag_dictionary": (
        "Tag Dictionary",
        "Complete canonical reference for every tag in scope.",
    ),
    "change_log": (
        "Change Log",
        "Pending Proficy export queue plus tags that need sync attention.",
    ),
    "commissioning": (
        "Commissioning Checklist",
        "Per-vessel readiness summary and verification table.",
    ),
    "network_map": (
        "Network / Address Map",
        "Tags grouped by PLC address family (%R, %G, %M, %AI, …).",
    ),
    "operator_manual": (
        "Operator Manual",
        "Narrative overview and vessel sections for operations staff.",
    ),
}


from models.documentation import (
    DocumentationColumn,
    DocumentationPackageResult,
    DocumentationTable,
)
class DocumentationService:
    """Builds documentation tables and writes multi-format export packages."""

    def __init__(
        self,
        sync_status_label: Callable[[str], str] | None = None,
    ) -> None:
        self._sync_label = sync_status_label or (
            lambda status: SYNC_STATUS_LABELS.get(
                status, status.replace("_", " ").title()
            )
        )

    def build_tables(
        self,
        tags: dict[str, TagRecord],
        *,
        selected_types: Iterable[str],
        vessel_filter: str | None = None,
        pending_exports: list[PendingExportChange] | None = None,
    ) -> list[DocumentationTable]:
        """Returns ordered tables for the requested documentation types."""
        scoped = self._filter_tags(tags, vessel_filter)
        tables: list[DocumentationTable] = []
        for doc_id in selected_types:
            if doc_id not in DOCUMENT_TYPES:
                continue
            builder = getattr(self, f"_build_{doc_id}", None)
            if builder is None:
                continue
            table = builder(scoped, vessel_filter=vessel_filter, pending_exports=pending_exports)
            if table is not None:
                tables.append(table)
        return tables

    def write_package(
        self,
        tables: list[DocumentationTable],
        output_dir: Path,
        *,
        write_html: bool = True,
        write_excel: bool = True,
        write_csv: bool = True,
        write_word: bool = False,
        vessel_filter: str | None = None,
        tag_count: int = 0,
    ) -> DocumentationPackageResult:
        """Writes selected formats into output_dir."""
        output_dir.mkdir(parents=True, exist_ok=True)
        result = DocumentationPackageResult(output_dir=output_dir)
        stamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

        if write_html:
            index_path = self._write_html_dashboard(
                tables,
                output_dir,
                generated_at=stamp,
                vessel_filter=vessel_filter,
                tag_count=tag_count,
            )
            result.written_files.append(index_path)
            for table in tables:
                page = output_dir / f"{table.doc_id}.html"
                self._write_html_page(table, page, generated_at=stamp)
                result.written_files.append(page)

        if write_excel and tables:
            excel_path = output_dir / "TagCentral_Documentation.xlsx"
            self._write_excel_workbook(tables, excel_path)
            result.written_files.append(excel_path)

        if write_csv:
            csv_dir = output_dir / "csv"
            csv_dir.mkdir(parents=True, exist_ok=True)
            for table in tables:
                path = csv_dir / f"{table.doc_id}.csv"
                self._write_csv_table(table, path)
                result.written_files.append(path)

        if write_word and tables:
            word_path = output_dir / "TagCentral_Documentation.docx"
            if self._write_word_document(tables, word_path, generated_at=stamp):
                result.written_files.append(word_path)

        return result

    def _filter_tags(
        self,
        tags: dict[str, TagRecord],
        vessel_filter: str | None,
    ) -> list[TagRecord]:
        if not vessel_filter or vessel_filter.upper() == "ALL":
            return sorted(tags.values(), key=lambda record: record.tag_name)
        vessel = vessel_filter.strip().upper()
        return sorted(
            (
                record
                for record in tags.values()
                if vessel in record.vessels
            ),
            key=lambda record: record.tag_name,
        )

    def _base_row(self, record: TagRecord) -> dict[str, str]:
        address = self._record_address(record)
        return {
            "tag_name": record.tag_name,
            "proficy_name": record.proficy_name or "",
            "cimplicity_pt_id": record.cimplicity_pt_id or "",
            "description": record.description,
            "address": address,
            "address_family": self._address_family(address),
            "vessels": ", ".join(sorted(record.vessels)),
            "sync_status": self._sync_label(record.sync_status),
            "link_method": record.link_method or "",
        }

    @staticmethod
    def _record_address(record: TagRecord) -> str:
        if record.linked_address:
            return record.linked_address
        return normalize_address(TagRecord._address_from_row(record.proficy_row_data))

    @staticmethod
    def _address_family(address: str) -> str:
        if not address or not is_resolvable_address(address):
            return "UNASSIGNED"
        match = _ADDRESS_FAMILY_PATTERN.match(address.strip().upper())
        if match:
            return f"%{match.group(1).upper()}"
        if address.upper().startswith("%AI"):
            return "%AI"
        return "OTHER"

    @staticmethod
    def _is_alarm_tag(record: TagRecord) -> bool:
        text = f"{record.tag_name} {record.description} {record.proficy_name}"
        if _ALARM_PATTERN.search(text):
            return True
        return False

    def _build_io_list(
        self,
        tags: list[TagRecord],
        **_: object,
    ) -> DocumentationTable:
        rows = []
        for record in tags:
            row = self._base_row(record)
            if row["address"] and is_resolvable_address(row["address"]):
                rows.append(row)
        return DocumentationTable(
            doc_id="io_list",
            title=DOCUMENT_TYPES["io_list"][0],
            summary=DOCUMENT_TYPES["io_list"][1],
            columns=[
                DocumentationColumn("tag_name", "Tag"),
                DocumentationColumn("description", "Description"),
                DocumentationColumn("address", "Address"),
                DocumentationColumn("address_family", "Family"),
                DocumentationColumn("proficy_name", "Proficy Name"),
                DocumentationColumn("cimplicity_pt_id", "Cimplicity PT_ID"),
                DocumentationColumn("vessels", "Vessels"),
                DocumentationColumn("sync_status", "Sync"),
            ],
            rows=rows,
        )

    def _build_alarm_list(
        self,
        tags: list[TagRecord],
        **_: object,
    ) -> DocumentationTable:
        rows = [self._base_row(record) for record in tags if self._is_alarm_tag(record)]
        return DocumentationTable(
            doc_id="alarm_list",
            title=DOCUMENT_TYPES["alarm_list"][0],
            summary=DOCUMENT_TYPES["alarm_list"][1],
            columns=[
                DocumentationColumn("tag_name", "Tag"),
                DocumentationColumn("description", "Description"),
                DocumentationColumn("address", "Address"),
                DocumentationColumn("proficy_name", "Proficy Name"),
                DocumentationColumn("cimplicity_pt_id", "Cimplicity PT_ID"),
                DocumentationColumn("sync_status", "Sync"),
                DocumentationColumn("vessels", "Vessels"),
            ],
            rows=rows,
        )

    def _build_tag_dictionary(
        self,
        tags: list[TagRecord],
        **_: object,
    ) -> DocumentationTable:
        rows = [self._base_row(record) for record in tags]
        return DocumentationTable(
            doc_id="tag_dictionary",
            title=DOCUMENT_TYPES["tag_dictionary"][0],
            summary=DOCUMENT_TYPES["tag_dictionary"][1],
            columns=[
                DocumentationColumn("tag_name", "Canonical Tag"),
                DocumentationColumn("description", "Description"),
                DocumentationColumn("address", "Address"),
                DocumentationColumn("proficy_name", "Proficy Name"),
                DocumentationColumn("cimplicity_pt_id", "Cimplicity PT_ID"),
                DocumentationColumn("sync_status", "Sync Status"),
                DocumentationColumn("link_method", "Link Method"),
                DocumentationColumn("vessels", "Vessels"),
            ],
            rows=rows,
        )

    def _build_change_log(
        self,
        tags: list[TagRecord],
        *,
        pending_exports: list[PendingExportChange] | None = None,
        **_: object,
    ) -> DocumentationTable:
        rows: list[dict[str, str]] = []
        for entry in pending_exports or []:
            name = str(entry.row_data.get("Name", "")).strip().upper()
            rows.append(
                {
                    "source": "Pending Proficy Export",
                    "tag_name": name,
                    "description": str(entry.row_data.get("Description", "")).strip().upper(),
                    "address": str(
                        entry.row_data.get("IOAddress", entry.row_data.get("Address", ""))
                    ).strip().upper(),
                    "vessels": entry.vessel,
                    "detail": ", ".join(
                        changed_field_labels(entry.baseline, entry.row_data)
                    ),
                }
            )
        for record in tags:
            if record.sync_status == SYNC_SYNCED:
                continue
            base = self._base_row(record)
            rows.append(
                {
                    "source": "Sync Review",
                    "tag_name": base["tag_name"],
                    "description": base["description"],
                    "address": base["address"],
                    "vessels": base["vessels"],
                    "detail": base["sync_status"],
                }
            )
        return DocumentationTable(
            doc_id="change_log",
            title=DOCUMENT_TYPES["change_log"][0],
            summary=DOCUMENT_TYPES["change_log"][1],
            columns=[
                DocumentationColumn("source", "Source"),
                DocumentationColumn("tag_name", "Tag"),
                DocumentationColumn("description", "Description"),
                DocumentationColumn("address", "Address"),
                DocumentationColumn("vessels", "Vessels"),
                DocumentationColumn("detail", "Details"),
            ],
            rows=rows,
        )

    def _build_commissioning(
        self,
        tags: list[TagRecord],
        *,
        vessel_filter: str | None = None,
        **_: object,
    ) -> DocumentationTable:
        by_vessel: dict[str, list[TagRecord]] = defaultdict(list)
        for record in tags:
            targets = sorted(record.vessels) or ["UNASSIGNED"]
            for vessel in targets:
                by_vessel[vessel].append(record)

        rows: list[dict[str, str]] = []
        for vessel in sorted(by_vessel):
            vessel_tags = by_vessel[vessel]
            linked = sum(1 for record in vessel_tags if record.cimplicity_pt_id)
            drift = sum(
                1
                for record in vessel_tags
                if record.sync_status
                in {SYNC_PROFICY_DRIFT, SYNC_NEEDS_ALIGN, SYNC_NAME_MISMATCH}
            )
            proficy_only = sum(
                1 for record in vessel_tags if record.sync_status == SYNC_PROFICY_ONLY
            )
            rows.append(
                {
                    "vessel": vessel,
                    "tag_count": str(len(vessel_tags)),
                    "cimplicity_linked": str(linked),
                    "needs_attention": str(drift),
                    "proficy_only": str(proficy_only),
                    "checklist": (
                        "Verify IO addresses in field | "
                        "Confirm Cimplicity PT_ID linkage | "
                        "Resolve sync drift before sail-away"
                    ),
                }
            )

        summary = DOCUMENT_TYPES["commissioning"][1]
        if vessel_filter and vessel_filter.upper() != "ALL":
            summary += f" Filtered to vessel {vessel_filter.upper()}."

        return DocumentationTable(
            doc_id="commissioning",
            title=DOCUMENT_TYPES["commissioning"][0],
            summary=summary,
            columns=[
                DocumentationColumn("vessel", "Vessel"),
                DocumentationColumn("tag_count", "Tags"),
                DocumentationColumn("cimplicity_linked", "Cimplicity Linked"),
                DocumentationColumn("needs_attention", "Needs Attention"),
                DocumentationColumn("proficy_only", "Proficy Only"),
                DocumentationColumn("checklist", "Commissioning Checklist"),
            ],
            rows=rows,
        )

    def _build_network_map(
        self,
        tags: list[TagRecord],
        **_: object,
    ) -> DocumentationTable:
        rows = []
        for record in tags:
            base = self._base_row(record)
            if not base["address"] or not is_resolvable_address(base["address"]):
                continue
            rows.append(base)
        rows.sort(key=lambda row: (row["address_family"], row["address"], row["tag_name"]))
        return DocumentationTable(
            doc_id="network_map",
            title=DOCUMENT_TYPES["network_map"][0],
            summary=DOCUMENT_TYPES["network_map"][1],
            columns=[
                DocumentationColumn("address_family", "Address Family"),
                DocumentationColumn("address", "Address"),
                DocumentationColumn("tag_name", "Tag"),
                DocumentationColumn("description", "Description"),
                DocumentationColumn("cimplicity_pt_id", "Cimplicity PT_ID"),
                DocumentationColumn("vessels", "Vessels"),
            ],
            rows=rows,
        )

    def _build_operator_manual(
        self,
        tags: list[TagRecord],
        **_: object,
    ) -> DocumentationTable:
        """Table of vessel chapters; HTML narrative is expanded in the page writer."""
        by_vessel: dict[str, int] = Counter()
        for record in tags:
            for vessel in record.vessels or {"UNASSIGNED"}:
                by_vessel[vessel] += 1
        rows = [
            {
                "section": vessel,
                "tag_count": str(count),
                "notes": (
                    "Operational tags for this vessel. See the HTML Operator Manual "
                    "page for narrative guidance and sync status overview."
                ),
            }
            for vessel, count in sorted(by_vessel.items())
        ]
        if not rows:
            rows.append(
                {
                    "section": "Overview",
                    "tag_count": "0",
                    "notes": "No tags in the current scope.",
                }
            )
        return DocumentationTable(
            doc_id="operator_manual",
            title=DOCUMENT_TYPES["operator_manual"][0],
            summary=DOCUMENT_TYPES["operator_manual"][1],
            columns=[
                DocumentationColumn("section", "Section"),
                DocumentationColumn("tag_count", "Tags"),
                DocumentationColumn("notes", "Notes"),
            ],
            rows=rows,
        )

    def _write_csv_table(self, table: DocumentationTable, path: Path) -> None:
        with path.open("w", newline="", encoding="utf-8") as file:
            writer = csv.DictWriter(
                file,
                fieldnames=[column.key for column in table.columns],
            )
            writer.writeheader()
            for row in table.rows:
                writer.writerow(
                    {column.key: row.get(column.key, "") for column in table.columns}
                )

    def _write_excel_workbook(self, tables: list[DocumentationTable], path: Path) -> None:
        pd = get_pandas()
        with pd.ExcelWriter(path, engine="openpyxl") as writer:
            for table in tables:
                sheet_name = table.doc_id[:31]
                frame = pd.DataFrame(
                    [
                        {column.header: row.get(column.key, "") for column in table.columns}
                        for row in table.rows
                    ]
                )
                frame.to_excel(writer, sheet_name=sheet_name, index=False)

    def _write_word_document(
        self,
        tables: list[DocumentationTable],
        path: Path,
        *,
        generated_at: str,
    ) -> bool:
        try:
            from docx import Document
        except ImportError:
            return False

        document = Document()
        document.add_heading("Tag Central — Engineering Documentation", level=0)
        document.add_paragraph(f"Generated {generated_at}")
        for table in tables:
            document.add_heading(table.title, level=1)
            document.add_paragraph(table.summary)
            if not table.rows:
                document.add_paragraph("(No rows in current scope.)")
                continue
            doc_table = document.add_table(
                rows=1,
                cols=len(table.columns),
                style="Table Grid",
            )
            doc_table.autofit = True
            header_cells = doc_table.rows[0].cells
            for index, column in enumerate(table.columns):
                header_cells[index].text = column.header
            for row in table.rows[:500]:
                cells = doc_table.add_row().cells
                for index, column in enumerate(table.columns):
                    cells[index].text = str(row.get(column.key, ""))
            if len(table.rows) > 500:
                document.add_paragraph(
                    f"(Showing first 500 of {len(table.rows)} rows. "
                    "Use Excel or CSV export for the full set.)"
                )
        document.save(path)
        return True

    def _write_html_dashboard(
        self,
        tables: list[DocumentationTable],
        output_dir: Path,
        *,
        generated_at: str,
        vessel_filter: str | None,
        tag_count: int,
    ) -> Path:
        index_path = output_dir / "index.html"
        vessel_text = vessel_filter if vessel_filter and vessel_filter != "ALL" else "All vessels"
        cards = []
        for table in tables:
            row_count = len(table.rows)
            cards.append(
                f"""
                <a class="card" href="{html.escape(table.doc_id)}.html">
                  <h2>{html.escape(table.title)}</h2>
                  <p>{html.escape(table.summary)}</p>
                  <span class="badge">{row_count} rows</span>
                </a>
                """
            )
        body = f"""
        <h1>Tag Central Documentation</h1>
        <p class="meta">Generated {html.escape(generated_at)} · {html.escape(vessel_text)} · {tag_count} tags in database</p>
        <p>Open any section below. Use your browser&rsquo;s <strong>Print → Save as PDF</strong> for PDF copies.</p>
        <div class="grid">{''.join(cards)}</div>
        """
        index_path.write_text(
            self._html_shell("Tag Central Documentation", body),
            encoding="utf-8",
        )
        return index_path

    def _write_html_page(
        self,
        table: DocumentationTable,
        path: Path,
        *,
        generated_at: str,
    ) -> None:
        extra = ""
        if table.doc_id == "operator_manual":
            extra = self._operator_manual_narrative(table)
        elif table.doc_id == "network_map":
            extra = self._network_map_summary(table)

        table_html = self._table_html(table)
        body = f"""
        <p><a href="index.html">&larr; Documentation home</a></p>
        <h1>{html.escape(table.title)}</h1>
        <p class="meta">{html.escape(table.summary)} · {len(table.rows)} rows · {html.escape(generated_at)}</p>
        {extra}
        {table_html}
        """
        path.write_text(self._html_shell(table.title, body), encoding="utf-8")

    def _operator_manual_narrative(self, table: DocumentationTable) -> str:
        parts = [
            "<section class='narrative'>",
            "<h2>How to use this manual</h2>",
            "<p>Each section lists tags assigned to a vessel. "
            "<strong>Synced</strong> tags match Cimplicity; "
            "<strong>Proficy Drift</strong> means the HMI and Proficy descriptions differ; "
            "resolve drift before operations rely on the tag.</p>",
            "<ul>",
        ]
        for row in table.rows:
            parts.append(
                f"<li><strong>{html.escape(row['section'])}</strong> — "
                f"{html.escape(row['tag_count'])} tags</li>"
            )
        parts.append("</ul></section>")
        return "".join(parts)

    def _network_map_summary(self, table: DocumentationTable) -> str:
        counts = Counter(row.get("address_family", "OTHER") for row in table.rows)
        items = "".join(
            f"<li><strong>{html.escape(family)}</strong>: {count} tags</li>"
            for family, count in sorted(counts.items())
        )
        return f"<section class='narrative'><h2>Address families</h2><ul>{items}</ul></section>"

    def _table_html(self, table: DocumentationTable) -> str:
        if not table.rows:
            return "<p><em>No data in the current filter scope.</em></p>"
        headers = "".join(
            f"<th>{html.escape(column.header)}</th>" for column in table.columns
        )
        body_rows = []
        for row in table.rows:
            cells = "".join(
                f"<td>{html.escape(str(row.get(column.key, '')))}</td>"
                for column in table.columns
            )
            body_rows.append(f"<tr>{cells}</tr>")
        return (
            "<div class='table-wrap'><table>"
            f"<thead><tr>{headers}</tr></thead>"
            f"<tbody>{''.join(body_rows)}</tbody></table></div>"
        )

    @staticmethod
    def _html_shell(title: str, body: str) -> str:
        return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1"/>
  <title>{html.escape(title)}</title>
  <style>
    :root {{
      --teal: #0d6b7a;
      --bg: #f4f7f8;
      --card: #ffffff;
      --text: #1f2937;
      --muted: #6b7280;
    }}
    body {{
      font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
      margin: 0; padding: 24px 32px 48px; background: var(--bg); color: var(--text);
      line-height: 1.45;
    }}
    h1 {{ color: var(--teal); margin-bottom: 0.25rem; }}
    h2 {{ color: var(--teal); font-size: 1.15rem; }}
    .meta {{ color: var(--muted); margin-top: 0; }}
    a {{ color: var(--teal); }}
    .grid {{
      display: grid; grid-template-columns: repeat(auto-fill, minmax(260px, 1fr));
      gap: 16px; margin-top: 24px;
    }}
    .card {{
      background: var(--card); border-radius: 10px; padding: 16px 18px;
      text-decoration: none; color: inherit; box-shadow: 0 1px 4px rgba(0,0,0,0.08);
      border: 1px solid #e5e7eb;
    }}
    .card:hover {{ border-color: var(--teal); }}
    .badge {{
      display: inline-block; background: #e0f2f4; color: var(--teal);
      padding: 2px 10px; border-radius: 999px; font-size: 0.85rem;
    }}
    .table-wrap {{ overflow-x: auto; background: var(--card); border-radius: 8px;
      padding: 8px; box-shadow: 0 1px 3px rgba(0,0,0,0.06); }}
    table {{ border-collapse: collapse; width: 100%; font-size: 0.9rem; }}
    th, td {{ border: 1px solid #e5e7eb; padding: 8px 10px; text-align: left; }}
    th {{ background: #e8f4f6; color: var(--teal); position: sticky; top: 0; }}
    tr:nth-child(even) {{ background: #fafbfc; }}
    .narrative {{ background: var(--card); padding: 16px 20px; border-radius: 8px;
      margin: 16px 0; border-left: 4px solid var(--teal); }}
    @media print {{
      body {{ background: white; }}
      .card {{ break-inside: avoid; }}
      th {{ position: static; }}
    }}
  </style>
</head>
<body>
{body}
</body>
</html>
"""
